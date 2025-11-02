"""Document loader for various formats"""
import os
import hashlib
import logging
from pathlib import Path
from typing import List, Optional
import mimetypes

from app.schemas.documents import Document

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Universal document loader supporting multiple formats"""
    
    def __init__(self, base_path: str = "./documents"):
        self.base_path = Path(base_path)
        self.processors = {
            'application/pdf': self._load_pdf,
            'text/plain': self._load_text,
            'text/markdown': self._load_text,
            'text/html': self._load_html,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._load_docx,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._load_excel,
            'text/csv': self._load_csv,
            'application/json': self._load_json,
        }
    
    async def load_all_documents(self) -> List[Document]:
        """Load all documents from base path"""
        documents = []
        
        for file_path in self.base_path.rglob('*'):
            if file_path.is_file():
                try:
                    doc = await self.load_single_document(str(file_path))
                    if doc:
                        documents.append(doc)
                except Exception as e:
                    logger.error(f"Failed to load {file_path}: {e}")
        
        return documents
    
    async def load_single_document(self, file_path: str) -> Optional[Document]:
        """Load a single document"""
        mime_type, _ = mimetypes.guess_type(file_path)
        
        if mime_type not in self.processors:
            logger.warning(f"Unsupported file type: {mime_type}")
            return None
        
        processor = self.processors[mime_type]
        return await processor(file_path)
    
    async def _load_pdf(self, file_path: str) -> Document:
        """Load PDF document"""
        from app.processing.pdf_processor import PDFProcessor
        processor = PDFProcessor()
        return await processor.extract(file_path)
    
    async def _load_text(self, file_path: str) -> Document:
        """Load plain text/markdown"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return Document(
            content=content,
            metadata={
                'filename': os.path.basename(file_path),
                'file_type': 'text',
                'size': os.path.getsize(file_path)
            },
            doc_id=self._generate_doc_id(file_path),
            source=file_path
        )
    
    async def _load_html(self, file_path: str) -> Document:
        """Load HTML document"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            logger.error("BeautifulSoup4 required for HTML processing")
            raise
        
        with open(file_path, 'r', encoding='utf-8') as f:
            html = f.read()
        
        soup = BeautifulSoup(html, 'html.parser')
        
        # Remove scripts and styles
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text(separator='\n', strip=True)
        
        return Document(
            content=text,
            metadata={
                'filename': os.path.basename(file_path),
                'file_type': 'html',
                'title': soup.title.string if soup.title else None
            },
            doc_id=self._generate_doc_id(file_path),
            source=file_path
        )
    
    async def _load_docx(self, file_path: str) -> Document:
        """Load Word document"""
        try:
            import docx
        except ImportError:
            logger.error("python-docx required for DOCX processing")
            raise
        
        doc = docx.Document(file_path)
        content = '\n'.join([para.text for para in doc.paragraphs])
        
        return Document(
            content=content,
            metadata={
                'filename': os.path.basename(file_path),
                'file_type': 'docx',
                'paragraphs': len(doc.paragraphs)
            },
            doc_id=self._generate_doc_id(file_path),
            source=file_path
        )
    
    async def _load_excel(self, file_path: str) -> Document:
        """Load Excel spreadsheet"""
        try:
            import pandas as pd
        except ImportError:
            logger.error("pandas required for Excel processing")
            raise
        
        # Read all sheets
        xl_file = pd.ExcelFile(file_path)
        content_parts = []
        
        for sheet_name in xl_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            content_parts.append(f"## Sheet: {sheet_name}\n")
            content_parts.append(df.to_markdown(index=False))
        
        return Document(
            content='\n\n'.join(content_parts),
            metadata={
                'filename': os.path.basename(file_path),
                'file_type': 'excel',
                'sheets': xl_file.sheet_names
            },
            doc_id=self._generate_doc_id(file_path),
            source=file_path
        )
    
    async def _load_csv(self, file_path: str) -> Document:
        """Load CSV file"""
        try:
            import pandas as pd
        except ImportError:
            logger.error("pandas required for CSV processing")
            raise
        
        df = pd.read_csv(file_path)
        content = df.to_markdown(index=False)
        
        return Document(
            content=content,
            metadata={
                'filename': os.path.basename(file_path),
                'file_type': 'csv',
                'rows': len(df),
                'columns': list(df.columns)
            },
            doc_id=self._generate_doc_id(file_path),
            source=file_path
        )
    
    async def _load_json(self, file_path: str) -> Document:
        """Load JSON file"""
        import json
        
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Convert JSON to readable text
        content = json.dumps(data, indent=2)
        
        return Document(
            content=content,
            metadata={
                'filename': os.path.basename(file_path),
                'file_type': 'json'
            },
            doc_id=self._generate_doc_id(file_path),
            source=file_path
        )
    
    def _generate_doc_id(self, file_path: str) -> str:
        """Generate unique document ID"""
        return hashlib.md5(file_path.encode()).hexdigest()

